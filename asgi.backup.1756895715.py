import os, glob, json, uuid, io
from typing import Dict, List, Any
from fastapi import FastAPI, HTTPException, Body
from fastapi.responses import HTMLResponse, FileResponse, StreamingResponse
from fastapi.staticfiles import StaticFiles
from pydantic import BaseModel
import openpyxl
from docx import Document
from docx.shared import Pt

DATA_DIR = os.path.join(os.getcwd(), "data")
AUDITS_DIR = os.path.join(DATA_DIR, "audits")
SESS_DIR = os.path.join(DATA_DIR, "sessions")

app = FastAPI(title="INFRATEC Audit Console")

# --------- Models ----------
class AskIn(BaseModel):
    question: str
    context: Dict[str, Any] = {}

# --------- In-memory stores ----------
PRESETS: Dict[str, Dict[str, List[Dict[str, str]]]] = {}  # iso -> section -> [{clause,question,file,row}]
SESSIONS: Dict[str, Dict[str, Any]] = {}  # id -> {header, entries}

def _infer_iso(fname: str) -> str:
    t = fname.lower()
    if "sector" in t and "8" in t: return "Sector Scheme 8"
    if "9001" in t: return "ISO 9001"
    if "14001" in t: return "ISO 14001"
    if "45001" in t: return "ISO 45001"
    if "27001" in t: return "ISO 27001"
    return "Unspecified"

def _clean(s):
    return (str(s).replace("\n"," ").strip() if s is not None else "")

def load_audits() -> Dict[str, Dict[str, List[Dict[str,str]]]]:
    presets: Dict[str, Dict[str, List[Dict[str,str]]]] = {}
    for path in glob.glob(os.path.join(AUDITS_DIR, "*.xlsx")):
        fname = os.path.basename(path)
        iso = _infer_iso(fname)
        try:
            wb = openpyxl.load_workbook(path, data_only=True)
        except Exception as e:
            print("[audits] cannot open", fname, e)
            continue
        for sheet in wb.sheetnames:
            ws = wb[sheet]
            section = sheet
            clause = ""
            for row in ws.iter_rows(values_only=True):
                cells = [_clean(c) for c in row if _clean(c)]
                if not cells: continue
                text = " ".join(cells)
                # crude: a row starting "5." etc becomes current clause
                if text[0:1].isdigit() and ("." in text[:3] or ":" in text[:3]):
                    clause = text
                    continue
                if not clause: 
                    continue
                presets.setdefault(iso, {}).setdefault(section, []).append({
                    "clause": clause,
                    "question": text,
                    "file": fname,
                    "row": str(getattr(row[0], "row", "?"))
                })
    return presets

def ensure_sessions_dir():
    os.makedirs(SESS_DIR, exist_ok=True)

# --------- Startup ----------
@app.on_event("startup")
def boot():
    global PRESETS
    ensure_sessions_dir()
    PRESETS = load_audits()
    print("[startup] loaded: ", {iso: sum(len(v) for v in sec.values()) for iso,sec in PRESETS.items()})

# --------- Health ----------
@app.get("/health")
def health(): return {"status":"ok"}

# --------- Presets / audits ----------
@app.post("/audits/_reload")
def reload_audits():
    global PRESETS
    PRESETS = load_audits()
    return {"ok": True, "isos": list(PRESETS.keys())}

@app.get("/audits")
def list_audits():
    return PRESETS

# --------- Session ----------
@app.post("/session/start")
def start_session(header: Dict[str, Any] = Body(default={})):
    sid = str(uuid.uuid4())[:8]
    SESSIONS[sid] = {"id": sid, "header": header, "entries": []}
    return SESSIONS[sid]

@app.post("/session/{sid}/save")
def save_entry(sid: str, entry: Dict[str, Any]):
    if sid not in SESSIONS: raise HTTPException(404, "session not found")
    SESSIONS[sid]["entries"].append(entry)
    return {"ok": True, "count": len(SESSIONS[sid]["entries"])}

# --------- Ask (stub that returns structured answer + cites) ----------
# NOTE: Replace this with your preferred RAG/LLM call. We keep it simple but
# include "sources" so they appear in UI and DOCX.
@app.post("/ask")
def ask(payload: AskIn):
    q = payload.question.strip()
    if not q: raise HTTPException(400, "empty")
    # minimal, rule-based cite finder from loaded audits:
    cites = []
    for iso, secmap in PRESETS.items():
        for sec, rows in secmap.items():
            for r in rows:
                if q.lower().split()[0] in r["question"].lower():
                    cites.append({"file": r["file"], "section": sec, "clause": r["clause"]})
                    if len(cites) >= 4: break
            if len(cites) >= 4: break
        if len(cites) >= 4: break
    answer = f"Draft answer for: {q}\n\n(Replace with LLM result.)"
    return {"answer": answer, "sources": cites}

# --------- Export DOCX (Cognito-style prep) ----------
@app.post("/export/cognito_prep")
def export_doc(payload: Dict[str, Any]):
    # payload: { header:{...}, entries:[{question, answer, sources:[]}] }
    header = payload.get("header", {})
    entries = payload.get("entries", [])

    doc = Document()
    doc.styles['Normal'].font.name = 'Calibri'
    doc.styles['Normal'].font.size = Pt(11)

    doc.add_heading("INFRATEC Audit – Cognito Prep", level=1)
    for k in ["Audit Number","Audit Date","Lead Auditor","Other Auditors",
              "Process to be audited","NHSS8 applicable?","Policies revised?",
              "Site per IMS Manual?","IMS Manual revised?"]:
        v = header.get(k, "")
        doc.add_paragraph(f"{k}: {v}")

    doc.add_heading("Findings", level=2)
    for i, e in enumerate(entries, 1):
        doc.add_paragraph(f"{i}. Question: {e.get('question','')}")
        doc.add_paragraph(f"   Answer: {e.get('answer','')}")
        if e.get("sources"):
            doc.add_paragraph("   Sources:")
            for s in e["sources"]:
                doc.add_paragraph(f"    • {s.get('file','')} — {s.get('section','')} — {s.get('clause','')}")

    buf = io.BytesIO()
    doc.save(buf); buf.seek(0)
    return StreamingResponse(buf, media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                             headers={"Content-Disposition": 'attachment; filename="cognito_prep.docx"'})

# --------- Static UI ----------
app.mount("/static", StaticFiles(directory="frontend"), name="static")

@app.get("/", response_class=HTMLResponse)
def index():
    with open(os.path.join("frontend","index.html"), "r", encoding="utf-8") as f:
        return HTMLResponse(f.read())
