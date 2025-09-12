
from pathlib import Path
import os
from fastapi import FastAPI
from ims_api import router as ims_router
from fastapi.responses import HTMLResponse, JSONResponse, PlainTextResponse
from fastapi.staticfiles import StaticFiles

app = FastAPI(title="InfraTec ISO Bot")


app.include_router(ims_router)
# === RUNTIME IMS INDEXER v2 ===
import os, threading, importlib

_IMS_RUNNING = False
_IMS_LAST_ERROR = None
_IMS_FILES_SEEN = 0

def _ims_worker():
    global _IMS_RUNNING, _IMS_LAST_ERROR, _IMS_FILES_SEEN
    _IMS_RUNNING = True
    _IMS_LAST_ERROR = None
    _IMS_FILES_SEEN = 0
    try:
        asgi_mod = importlib.import_module("asgi")
        base = getattr(asgi_mod, "IMS_DIR", os.path.join("data", "source_docs"))
        # Prefer extract_text_from_file, fallback to extract_text(path)
        extract = getattr(asgi_mod, "extract_text_from_file", None) or getattr(asgi_mod, "extract_text", None)
        if not extract:
            raise RuntimeError("No text extractor found in asgi.py")

        index = []
        for root, dirs, files in os.walk(base):
            if "__MACOSX" in root:
                continue
            for fn in files:
                _IMS_FILES_SEEN += 1
                ext = os.path.splitext(fn)[1].lower()
                if ext not in (".txt", ".md", ".pdf", ".docx", ".xlsx", ".xlsm", ".xltx", ".xltm"):
                    continue
                path = os.path.join(root, fn)
                try:
                    text = extract(path)
                except Exception:
                    text = ""
                if not text:
                    continue
                rel = os.path.relpath(path, base)
                # Chunk into ~1200 chars so we don't blow memory
                for i in range(0, len(text), 1200):
                    chunk = text[i:i+1200]
                    if chunk.strip():
                        index.append({"relpath": rel, "text": chunk})

        setattr(asgi_mod, "IMS_INDEX", index)
    except Exception as e:
        _IMS_LAST_ERROR = str(e)
    finally:
        _IMS_RUNNING = False

@app.post("/ims/reindex")
def ims_reindex():
    t = threading.Thread(target=_ims_worker, daemon=True)
    t.start()
    return {"ok": True, "started": True}

@app.get("/ims/status")
def ims_status():
    asgi_mod = importlib.import_module("asgi")
    idx = getattr(asgi_mod, "IMS_INDEX", None)
    chunks = len(idx) if isinstance(idx, list) else 0
    return {
        "running": _IMS_RUNNING,
        "last_error": _IMS_LAST_ERROR,
        "files_seen": _IMS_FILES_SEEN,
        "chunks": chunks
    }
# === END RUNTIME IMS INDEXER v2 ===


# ---- Try to import your main ASGI app + helpers from asgi.py ----
asgi_mod = None
try:
    import asgi as asgi_mod
    if hasattr(asgi_mod, "app"):
        app.mount("/core", asgi_mod.app)
        print("[server] Mounted asgi.app at /core")
except Exception as e:
    print("[server] asgi import failed:", e)

build_ims_index = getattr(asgi_mod, "build_ims_index", None) if asgi_mod else None
load_ims_index  = getattr(asgi_mod, "load_ims_index",  None) if asgi_mod else None
IMS_INDEX_PATH  = getattr(asgi_mod, "IMS_INDEX_PATH",  "data/ims_index.json") if asgi_mod else "data/ims_index.json"

ask_func = None
if asgi_mod:
    for _name in ("ask", "ask_ims", "answer_question", "answer", "run_query"):
        if hasattr(asgi_mod, _name):
            ask_func = getattr(asgi_mod, _name)
            break

# ---- Static / frontend ----
if Path("static").exists():
    app.mount("/static", StaticFiles(directory="static"), name="static")
if Path("frontend").exists():
    app.mount("/frontend", StaticFiles(directory="frontend"), name="frontend")

def _read_index_html() -> str:
    for p in (Path("static/index.html"), Path("frontend/index.html")):
        if p.exists():
            return p.read_text(encoding="utf-8")
    return "<h1>InfraTec ISO Bot</h1><p>Place your UI at static/index.html or frontend/index.html.</p>"

# ---- Routes ----
@app.get("/health")
def health():
    return {"status": "ok"}

@app.get("/", response_class=HTMLResponse)
def root():
    return HTMLResponse(_read_index_html())

@app.on_event("startup")
async def ensure_index():
    if not (build_ims_index and load_ims_index):
        print("[startup] IMS index builder not available (using runtime search only)")
        return
    try:
        needs = (not os.path.exists(IMS_INDEX_PATH)) or (os.path.getsize(IMS_INDEX_PATH) == 0)
        if needs:
            print("[startup] No IMS index — building…")
            n = build_ims_index()
            print(f"[startup] built {n} IMS chunks")
        else:
            print("[startup] IMS index present")
        load_ims_index()
    except Exception as e:
        print("[startup] IMS index error:", e)

@app.post("/admin/reindex")
def admin_reindex():
    if not (build_ims_index and load_ims_index):
        return {"status": "error", "error": "index builder not available"}
    n = build_ims_index()
    load_ims_index()
    return {"status": "ok", "chunks": n}

# Minimal /ask passthrough
from pydantic import BaseModel
class AskBody(BaseModel):
    question: str

@app.post("/ask")
def ask_endpoint(body: AskBody):
    if not ask_func:
        return {"answer": "[server] ask() is not wired from asgi.py", "sources": []}
    try:
        res = ask_func(body.question)
        if isinstance(res, dict):
            return res
        return {"answer": str(res), "sources": []}
    except Exception as e:
        return {"answer": f"[LLM error: {e}]", "sources": []}


@app.get("/_debug/ims", response_class=JSONResponse)
async def _debug_ims():
    import os
    base = os.getenv("IMS_DIR") or os.path.join(os.path.dirname(__file__), "data", "source_docs")
    files = []
    for root, dirs, filenames in os.walk(base):
        # skip macOS artifact folders / files
        dirs[:] = [d for d in dirs if not d.startswith("__MACOSX")]
        for name in filenames:
            if name.startswith("._"):
                continue
            files.append(os.path.relpath(os.path.join(root, name), base))
    return JSONResponse({"ims_dir": base, "files_found": len(files), "sample": files[:20]})


# === RUNTIME IMS INDEXER ===
# Lightweight on-demand indexer over data/source_docs
import os, pathlib
from typing import List, Dict

try:
    app
except NameError:
    from fastapi import FastAPI
    app = FastAPI()

app.include_router(ims_router)

IMS_DIR = os.getenv("IMS_DIR", os.path.join(os.path.dirname(__file__), "data", "source_docs"))
IMS_INDEX: List[Dict] = []

def _iter_files(root):
    exts = {".txt", ".md", ".pdf", ".docx", ".xlsx", ".xlsm"}
    r = pathlib.Path(root)
    for p in r.rglob("*"):
        if p.is_file() and p.suffix.lower() in exts and "/__MACOSX/" not in str(p):
            yield p

def _read_file(path: pathlib.Path) -> str:
    ext = path.suffix.lower()
    try:
        if ext in (".txt", ".md"):
            return path.read_text(encoding="utf-8", errors="ignore")
        if ext == ".pdf":
            try:
                from pdfminer.high_level import extract_text as _pdf_extract
                return _pdf_extract(str(path)) or ""
            except Exception:
                return ""
        if ext == ".docx":
            try:
                from docx import Document
                return "\n".join((p.text or "").strip() for p in Document(str(path)).paragraphs)
            except Exception:
                return ""
        if ext in (".xlsx", ".xlsm"):
            try:
                import openpyxl
                wb = openpyxl.load_workbook(filename=str(path), data_only=True, read_only=True)
                out=[]; remaining=500
                for ws in wb.worksheets[:2]:
                    out.append(f"# {ws.title}")
                    for row in ws.iter_rows(values_only=True):
                        line = " ".join("" if v is None else str(v) for v in row).strip()
                        if line:
                            out.append(line)
                            remaining -= 1
                            if remaining <= 0:
                                break
                return "\n".join(out)
            except Exception:
                return ""
    except Exception:
        return ""
    return ""

def _chunk(text: str, size: int = 1200, overlap: int = 200):
    i, n = 0, len(text)
    while i < n:
        yield text[i:i+size]
        i += max(1, size - overlap)

def rebuild_index():
    IMS_INDEX.clear()
    count = 0
    root = pathlib.Path(IMS_DIR)
    for p in _iter_files(root):
        txt = _read_file(p)
        if not txt:
            continue
        rel = str(p.relative_to(root))
        for ch in _chunk(txt):
            IMS_INDEX.append({"file": rel, "text": ch})
            count += 1
    return count

from fastapi import APIRouter
router_runtime = APIRouter()

@router_runtime.post("/admin/reindex")
async def _admin_reindex():
    try:
        cnt = rebuild_index()
        return {"status": "ok", "chunks": cnt}
    except Exception as e:
        return {"status": "error", "error": str(e)}

@router_runtime.get("/_debug/ims")
async def _debug_ims():
    try:
        sample = []
        for i, p in zip(range(20), _iter_files(IMS_DIR)):
            sample.append(str(p.relative_to(IMS_DIR)))
        total = sum(1 for _ in _iter_files(IMS_DIR))
        return {"ims_dir": IMS_DIR, "files_found": total, "sample": sample}
    except Exception as e:
        return {"ims_dir": IMS_DIR, "files_found": 0, "error": str(e)}

app.include_router(router_runtime)
# === END RUNTIME IMS INDEXER ===


# === RUNTIME IMS INDEXER (force) ===
from fastapi import Body
import os, json
from pathlib import Path as _Path
from io import BytesIO
IMS_INDEX = []

def _ims_dir():
    try:
        return IMS_DIR
    except NameError:
        base = os.path.join(os.path.dirname(__file__), "data", "source_docs")
        return os.getenv("IMS_DIR", base)

def _extract_text_generic(path):
    ext = _Path(path).suffix.lower()
    try:
        if ext in {".txt",".md"}:
            return _Path(path).read_text(encoding="utf-8", errors="ignore")
        if ext == ".docx":
            try:
                from docx import Document
                doc = Document(path)
                return "\n".join([(p.text or "").strip() for p in doc.paragraphs if (p.text or "").strip()])
            except Exception:
                return ""
        if ext in {".xlsx",".xlsm"}:
            try:
                import openpyxl
                wb = openpyxl.load_workbook(filename=path, data_only=True, read_only=True)
                out=[]
                for ws in wb.worksheets[:3]:
                    out.append(f"# SHEET: {ws.title}")
                    n=0
                    for row in ws.iter_rows(values_only=True):
                        line=" ".join(str(v) for v in row if v is not None).strip()
                        if line:
                            out.append(line); n+=1
                            if n>=500: break
                return "\n".join(out)
            except Exception:
                return ""
        if ext == ".pdf":
            try:
                from pdfminer.high_level import extract_text
                return extract_text(path) or ""
            except Exception:
                return ""
    except Exception:
        return ""
    return ""

def _chunk(text, maxlen=1000):
    text = " ".join(text.split())
    chunks=[]
    i=0
    while i < len(text):
        j=min(len(text), i+maxlen)
        chunks.append(text[i:j])
        i=j
    return [c for c in chunks if c]

@app.post("/admin/reindex_force")
def admin_reindex_force():
    global IMS_INDEX
    ims_dir = _ims_dir()
    files=[]
    for root, dirs, fnames in os.walk(ims_dir):
        for fn in fnames:
            if fn.startswith("._"):
                continue
            if any(fn.lower().endswith(ext) for ext in (".txt",".md",".pdf",".docx",".xlsx",".xlsm")):
                files.append(os.path.join(root, fn))
    chunks=[]
    for fp in files:
        txt = _extract_text_generic(fp)
        if not txt:
            continue
        for c in _chunk(txt, 1000):
            chunks.append({"file": os.path.relpath(fp, ims_dir), "text": c})
    IMS_INDEX = chunks
    return {"status":"ok","files":len(files),"chunks":len(chunks)}

@app.get("/_debug/index")
def debug_index():
    return {"chunks": len(IMS_INDEX)}
# === END RUNTIME IMS INDEXER (force) ===


# === BEGIN: Simple runtime IMS indexer (self-contained) ===
from fastapi import BackgroundTasks

def _ims_build_index():
    """Walk IMS_DIR, extract text, chunk it, and store in both this module and asgi."""
    import os, sys
    try:
        import asgi as _asgi
    except Exception:
        _asgi = None

    # Resolve root
    root = None
    if _asgi and hasattr(_asgi, "IMS_DIR"):
        root = getattr(_asgi, "IMS_DIR")
    if not root:
        root = os.path.join(os.getcwd(), "data", "source_docs")

    # Lazy imports inside the function to avoid import-time errors
    def _extract_any(fp: str) -> str:
        ext = os.path.splitext(fp)[1].lower()
        # Prefer asgi.extract_text if present
        if _asgi and hasattr(_asgi, "extract_text"):
            try:
                return _asgi.extract_text(fp) or ""
            except Exception:
                pass
        # Fallbacks
        if ext in (".txt", ".md"):
            try:
                with open(fp, "r", encoding="utf-8", errors="ignore") as f:
                    return f.read()
            except Exception:
                return ""
        if ext == ".docx":
            try:
                from docx import Document
                d = Document(fp)
                return "\n".join(p.text for p in d.paragraphs if (p.text or "").strip())
            except Exception:
                return ""
        if ext == ".pdf":
            try:
                from PyPDF2 import PdfReader
                pdf = PdfReader(fp)
                return "\n".join((page.extract_text() or "") for page in getattr(pdf, "pages", []))
            except Exception:
                return ""
        if ext in (".xlsx", ".xlsm", ".xltx", ".xltm"):
            try:
                import openpyxl
                wb = openpyxl.load_workbook(fp, data_only=True, read_only=True)
                out = []
                for ws in wb.worksheets[:3]:
                    out.append(f"# SHEET: {ws.title}")
                    count = 0
                    for row in ws.iter_rows(values_only=True):
                        line = " ".join(str(v) for v in row if v is not None).strip()
                        if line:
                            out.append(line)
                            count += 1
                            if count >= 500:
                                break
                return "\n".join(out)
            except Exception:
                return ""
        return ""

    # Walk and index
    exts = (".txt", ".md", ".docx", ".pdf", ".xlsx", ".xlsm", ".xltx", ".xltm")
    index = []
    seen_files = set()
    for dirpath, _, files in os.walk(root):
        for name in files:
            if name.startswith("._"):
                continue
            if not name.lower().endswith(exts):
                continue
            full = os.path.join(dirpath, name)
            rel = os.path.relpath(full, root)
            txt = _extract_any(full)
            if not txt:
                continue
            seen_files.add(rel)
            # simple overlapping chunker
            chunk_size, step = 1200, 1000
            for i in range(0, len(txt), step):
                chunk = txt[i:i+chunk_size]
                if chunk.strip():
                    index.append({"relpath": rel, "text": chunk})

    # Store in this module (server) and also in asgi if available
    import sys as _sys
    _this = _sys.modules[__name__]
    setattr(_this, "IMS_INDEX", index)
    if _asgi is not None:
        try:
            setattr(_asgi, "IMS_INDEX", index)
        except Exception:
            pass
    return {"files": len(seen_files), "chunks": len(index), "root": root}

@app.post("/ims/reindex")
def ims_reindex(tasks: BackgroundTasks):
    tasks.add_task(_ims_build_index)
    return {"ok": True, "started": True}

@app.get("/_debug/index2")
def debug_index2():
    idx = globals().get("IMS_INDEX")
    return {"chunks": 0 if idx is None else len(idx)}
# === END: Simple runtime IMS indexer (self-contained) ===


# === LIGHTWEIGHT RUNTIME IMS INDEXER ===
import os, threading, importlib

# Indexer runtime state (for /ims/status)
_IMS_RUNNING = False
_IMS_LAST_ERROR = None
_IMS_FILES_SEEN = 0

def _ims_worker():
    global _IMS_RUNNING, _IMS_LAST_ERROR, _IMS_FILES_SEEN
    _IMS_RUNNING = True
    _IMS_LAST_ERROR = None
    _IMS_FILES_SEEN = 0
    try:
        asgi_mod = importlib.import_module("asgi")
        base = getattr(asgi_mod, "IMS_DIR", os.path.join("data", "source_docs"))
        extract = (
            getattr(asgi_mod, "extract_text_from_file", None)
            or getattr(asgi_mod, "extract_text", None)
        )
        if not extract:
            raise RuntimeError("no extract_text function in asgi")

        index = []
        for root, dirs, files in os.walk(base):
            # skip macOS metadata dirs
            if "__MACOSX" in root:
                continue
            for fn in files:
                _IMS_FILES_SEEN += 1
                ext = os.path.splitext(fn)[1].lower()
                if ext not in (".txt", ".md", ".pdf", ".docx", ".xlsx", ".xlsm", ".xltx", ".xltm"):
                    continue
                path = os.path.join(root, fn)
                try:
                    text = extract(path)
                except Exception:
                    text = ""
                if not text:
                    continue
                # simple chunking ~1200 chars per chunk to keep memory reasonable
                rel = os.path.relpath(path, base)
                for i in range(0, len(text), 1200):
                    chunk = text[i:i+1200]
                    if chunk.strip():
                        index.append({"relpath": rel, "text": chunk})
        # publish into asgi module so /ask can use it
        setattr(asgi_mod, "IMS_INDEX", index)
    except Exception as e:
        _IMS_LAST_ERROR = str(e)
    finally:
        _IMS_RUNNING = False

@app.post("/ims/reindex")
def ims_reindex():
    t = threading.Thread(target=_ims_worker, daemon=True)
    t.start()
    return {"ok": True, "started": True}

@app.get("/ims/status")
def ims_status():
    asgi_mod = importlib.import_module("asgi")
    idx = getattr(asgi_mod, "IMS_INDEX", None)
    chunks = len(idx) if isinstance(idx, list) else 0
    return {
        "running": _IMS_RUNNING,
        "last_error": _IMS_LAST_ERROR,
        "files_seen": _IMS_FILES_SEEN,
        "chunks": chunks,
    }
# === END LIGHTWEIGHT RUNTIME IMS INDEXER ===
