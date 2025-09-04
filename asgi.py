import os, json, io, uuid, math, glob
from typing import Dict, List, Any, Tuple
from fastapi import FastAPI, HTTPException, Body
from fastapi.responses import HTMLResponse, StreamingResponse
from fastapi.staticfiles import StaticFiles
from pydantic import BaseModel
import openpyxl
from docx import Document
from docx.shared import Pt
from openai import OpenAI

# Optional PDF support (if pypdf installed)
try:
    from pypdf import PdfReader
except Exception:
    PdfReader = None

DATA_DIR = os.path.join(os.getcwd(), "data")
AUDITS_DIR = os.path.join(DATA_DIR, "audits")
IMS_DIR = os.path.join(DATA_DIR, "source_docs", "ManagementSystem")
IMS_INDEX_PATH = os.path.join(DATA_DIR, "ims_index.json")
SESS_DIR = os.path.join(DATA_DIR, "sessions")

app = FastAPI(title="INFRATEC Audit Console")
client = OpenAI()

class AskIn(BaseModel):
    question: str
    context: Dict[str, Any] = {}

PRESETS: Dict[str, Dict[str, List[Dict[str, str]]]] = {}
IMS_INDEX: List[Dict[str, Any]] = []
SESSIONS: Dict[str, Dict[str, Any]] = {}

def _infer_iso(fname: str) -> str:
    t = fname.lower()
    if "sector" in t and "8" in t: return "Sector Scheme 8"
    if "9001" in t: return "ISO 9001"
    if "14001" in t: return "ISO 14001"
    if "45001" in t: return "ISO 45001"
    if "27001" in t: return "ISO 27001"
    return "Unspecified"

def _clean(s): return (str(s).replace("\n"," ").strip() if s is not None else "")

def load_audits() -> Dict[str, Dict[str, List[Dict[str,str]]]]:
    presets: Dict[str, Dict[str, List[Dict[str,str]]]] = {}
    for path in glob.glob(os.path.join(AUDITS_DIR, "*.xlsx")):
        fname = os.path.basename(path)
        iso = _infer_iso(fname)
        try:
            wb = openpyxl.load_workbook(path, data_only=True)
        except Exception as e:
            print("[audits] cannot open", fname, e); continue
        for sheet in wb.sheetnames:
            ws = wb[sheet]
            section = sheet
            clause = ""
            for row in ws.iter_rows(values_only=True):
                cells = [_clean(c) for c in row if _clean(c)]
                if not cells: continue
                text = " ".join(cells)
                if text[:1].isdigit() and ("." in text[:3] or ":" in text[:3]):
                    clause = text; continue
                if not clause: continue
                presets.setdefault(iso, {}).setdefault(section, []).append({
                    "clause": clause,
                    "question": text,
                    "file": fname,
                    "row": "?"
                })
    return presets

def read_txt(path:str) -> str:
    with open(path, "r", encoding="utf-8", errors="ignore") as f:
        return f.read()

def read_docx(path:str) -> str:
    d = Document(path)
    return "\n".join(p.text for p in d.paragraphs)

def read_pdf(path:str) -> str:
    if PdfReader is None: return ""
    try:
        pdf = PdfReader(path)
        return "\n".join([p.extract_text() or "" for p in pdf.pages])
    except Exception as e:
        print("[ims][pdf] failed", path, e)
        return ""

def extract_text(path:str) -> str:
    ext = os.path.splitext(path)[1].lower()
    if ext in [".txt", ".md"]: return read_txt(path)
    if ext == ".docx": return read_docx(path)
    if ext == ".pdf": return read_pdf(path)
    return ""

def chunk_text(text:str, max_chars:int=1200, overlap:int=150) -> List[str]:
    text = text.replace("\r","").strip()
    if not text: return []
    chunks, i = [], 0
    while i < len(text):
        j = min(len(text), i+max_chars)
        chunk = text[i:j].strip()
        if chunk: chunks.append(chunk)
        i = j - overlap
        if i < 0: i = 0
        if j >= len(text): break
    return chunks

def embed_texts(texts: List[str]) -> List[List[float]]:
    if not texts: return []
    resp = client.embeddings.create(model="text-embedding-3-large", input=texts)
    return [d.embedding for d in resp.data]

def _norm(v: List[float]) -> float:
    return math.sqrt(sum(x*x for x in v)) or 1.0

def _cos(a: List[float], b: List[float], nb: float=None) -> float:
    if nb is None: nb = _norm(b)
    na = _norm(a)
    dot = sum(x*y for x,y in zip(a,b))
    return dot / (na*nb)

def build_ims_index() -> List[Dict[str,Any]]:
    index: List[Dict[str,Any]] = []
    files = []
    if os.path.isdir(IMS_DIR):
        for root,_,fnames in os.walk(IMS_DIR):
            for fn in fnames:
                if os.path.splitext(fn)[1].lower() in [".txt",".md",".docx",".pdf"]:
                    files.append(os.path.join(root, fn))
    files.sort()
    for path in files:
        rel = os.path.relpath(path, IMS_DIR)
        text = extract_text(path)
        if not text: continue
        chunks = chunk_text(text)
        for i in range(0, len(chunks), 32):
            batch = chunks[i:i+32]
            embs = embed_texts(batch)
            for ch, em in zip(batch, embs):
                index.append({
                    "id": str(uuid.uuid4())[:8],
                    "relpath": rel,
                    "file": os.path.basename(path),
                    "chunk": ch,
                    "embedding": em,
                    "norm": _norm(em)
                })
    os.makedirs(DATA_DIR, exist_ok=True)
    with open(IMS_INDEX_PATH, "w", encoding="utf-8") as f:
        json.dump(index, f)
    return index

def load_ims_index() -> List[Dict[str,Any]]:
    if not os.path.exists(IMS_INDEX_PATH): return []
    with open(IMS_INDEX_PATH, "r", encoding="utf-8") as f:
        return json.load(f)

def ims_search(q: str, k: int=6) -> List[Dict[str,Any]]:
    if not IMS_INDEX: return []
    q_emb = embed_texts([q])[0]
    nq = _norm(q_emb)
    scored: List[Tuple[float,Dict[str,Any]]] = []
    for rec in IMS_INDEX:
        s = _cos(q_emb, rec["embedding"], rec.get("norm"))
        scored.append((s, rec))
    scored.sort(key=lambda x: x[0], reverse=True)
    return [r for _,r in scored[:k]]

def ensure_sessions_dir(): os.makedirs(SESS_DIR, exist_ok=True)

@app.on_event("startup")
def boot():
    global PRESETS, IMS_INDEX
    ensure_sessions_dir()
    PRESETS = load_audits()
    IMS_INDEX = load_ims_index()
    print("[startup] loaded:", {iso: sum(len(v) for v in sec.values()) for iso,sec in PRESETS.items()})
    print("[startup] ims chunks:", len(IMS_INDEX))

@app.get("/health")
def health(): return {"status":"ok"}

@app.post("/audits/_reload")
def reload_audits():
    global PRESETS
    PRESETS = load_audits()
    return {"ok": True, "isos": list(PRESETS.keys())}

@app.get("/audits")
def list_audits(): return PRESETS

# ---- IMS endpoints ----
@app.post("/ims/_reindex")
def ims_reindex():
    global IMS_INDEX
    IMS_INDEX = build_ims_index()
    return {"ok": True, "chunks": len(IMS_INDEX)}

@app.get("/ims/_debug")
def ims_debug():
    files = {}
    for rec in IMS_INDEX:
        files[rec["relpath"]] = files.get(rec["relpath"], 0) + 1
    return {"chunks": len(IMS_INDEX), "files": files}

# ---- Sessions ----
@app.post("/session/start")
def start_session(header: Dict[str, Any] = Body(default={})):
    sid = str(uuid.uuid4())[:8]
    SESSIONS[sid] = {"id": sid, "header": header, "entries": []}
    return SESSIONS[sid]

@app.post("/session/{sid}/save")
def save_entry(sid: str, entry: Dict[str, Any]):
    if sid not in SESSIONS: raise HTTPException(404, "session not found")
    SESSIONS[sid]["entries"].append(entry)
    return {"ok": True, "count": len(SESSIONS[sid]["entries"])}

# ---- Ask (uses audits + IMS if available) ----
@app.post("/ask")
def ask(payload: AskIn):
    q = payload.question.strip()
    if not q: raise HTTPException(400, "empty")

    # Audit hits
    hits = []
    tokens = set(q.lower().split())
    for iso, secmap in PRESETS.items():
        for sec, rows in secmap.items():
            for r in rows:
                txt = (r["clause"] + " " + r["question"]).lower()
                score = len(tokens.intersection(set(txt.split())))
                if score > 0:
                    hits.append((score, {"file": r["file"], "section": sec, "clause": r["clause"], "question": r["question"]}))
    hits.sort(key=lambda x: x[0], reverse=True)
    audit_top = [h[1] for h in hits[:6]]

    # IMS hits (if indexed)
    ims_top = ims_search(q, k=6)

    audit_ctx = "\n".join(f"- [AUDIT {h['file']} — {h['section']} — {h['clause']}] {h['question']}" for h in audit_top) or "No audit checklist matches."
    ims_ctx = "\n".join(f"- [IMS {h['relpath']}] {h['chunk'][:600]}" for h in ims_top) or "No IMS excerpts found."

    system_msg = (
        "You are INFRATEC's audit assistant. "
        "Answer concisely in UK English. Cite clauses and recommend evidence. "
        "Use IMS excerpts for company-specific practice; if gaps exist, say what evidence is required."
    )
    user_msg = f"Question:\n{q}\n\nRelevant audit checklist rows:\n{audit_ctx}\n\nRelevant IMS excerpts:\n{ims_ctx}\n\nWrite a practical answer for an internal audit note."

    try:
        resp = client.chat.completions.create(
            model="gpt-4.1-mini",
            messages=[{"role":"system","content":system_msg},{"role":"user","content":user_msg}],
            temperature=0.2,
        )
        answer = resp.choices[0].message.content.strip()
    except Exception as e:
        print("[ask] OpenAI error:", e)
        answer = f"Draft answer for: {q}\n\n(LLM call failed — {e})"

    sources = []
    for h in audit_top[:4]:
        sources.append({"file": h["file"], "section": h["section"], "clause": h["clause"]})
    for h in ims_top[:4]:
        sources.append({"file": h["file"], "section": "IMS", "clause": h["relpath"]})
    return {"answer": answer, "sources": sources}

# ---- Export ----
@app.post("/export/cognito_prep")
def export_doc(payload: Dict[str, Any]=Body(default={"header":{}, "entries":[]})):
    header = payload.get("header", {})
    entries = payload.get("entries", [])

    doc = Document()
    doc.styles['Normal'].font.name = 'Calibri'
    doc.styles['Normal'].font.size = Pt(11)

    doc.add_heading("INFRATEC Audit – Cognito Prep", level=1)
    for k in ["Audit Number","Audit Date","Lead Auditor","Other Auditors",
              "Process to be audited","NHSS8 applicable?","Policies revised?",
              "Site per IMS Manual?","IMS Manual revised?"]:
        doc.add_paragraph(f"{k}: {header.get(k,'')}")

    doc.add_heading("Findings", level=2)
    for i, e in enumerate(entries, 1):
        doc.add_paragraph(f"{i}. Question: {e.get('question','')}")
        doc.add_paragraph(f"   Answer: {e.get('answer','')}")
        if e.get("sources"):
            doc.add_paragraph("   Sources:")
            for s in e["sources"]:
                doc.add_paragraph(f"    • {s.get('file','')} — {s.get('section','')} — {s.get('clause','')}")

    buf = io.BytesIO(); doc.save(buf); buf.seek(0)
    return StreamingResponse(buf,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": 'attachment; filename="cognito_prep.docx"'})

# ---- Static UI ----
app.mount("/static", StaticFiles(directory="frontend"), name="static")

@app.get("/", response_class=HTMLResponse)
def index():
    with open(os.path.join("frontend","index.html"), "r", encoding="utf-8") as f:
        return HTMLResponse(f.read())

# ---- IMS bootstrap from URL (Google Drive / any direct URL) ----
def _has_any_ims_files() -> bool:
    if not os.path.isdir(IMS_DIR): return False
    for root,_,files in os.walk(IMS_DIR):
        for fn in files:
            if os.path.splitext(fn)[1].lower() in [".txt",".md",".docx",".pdf"]:
                return True
    return False

def bootstrap_ims_from_url():
    ims_url = os.getenv("IMS_URL", "").strip()
    force = os.getenv("IMS_FORCE_FETCH", "0").strip() in ("1","true","yes")
    if not ims_url:
        return
    if _has_any_ims_files() and not force:
        print("[ims][bootstrap] files already present; skip download")
        return
    try:
        os.makedirs(IMS_DIR, exist_ok=True)
        zip_path = "/tmp/ims.zip"
        try:
            import gdown
            print(f"[ims][bootstrap] downloading via gdown: {ims_url}")
            gdown.download(url=ims_url, output=zip_path, quiet=False, fuzzy=True)
        except Exception as e:
            print("[ims][bootstrap] gdown failed:", e, "trying curl -L")
            os.system(f'curl -L -o "{zip_path}" "{ims_url}"')

        # unzip
        print(f"[ims][bootstrap] unzipping to {IMS_DIR}")
        os.system(f'unzip -o "{zip_path}" -d "{IMS_DIR}" > /dev/null 2>&1 || true')
    except Exception as e:
        print("[ims][bootstrap] error:", e)

# call bootstrap BEFORE loading IMS index
try:
    bootstrap_ims_from_url()
except Exception as e:
    print("[ims][bootstrap] skipped due to error:", e)

# ===== Async IMS reindex (background) =====
import threading
IMS_REINDEXING = False
IMS_LAST_ERROR = None

def _reindex_worker():
    global IMS_INDEX, IMS_REINDEXING, IMS_LAST_ERROR
    try:
        IMS_LAST_ERROR = None
        IMS_INDEX = build_ims_index()
    except Exception as e:
        IMS_LAST_ERROR = str(e)
    finally:
        IMS_REINDEXING = False

@app.post("/ims/_reindex")
def ims_reindex_async():
    global IMS_REINDEXING
    if IMS_REINDEXING:
        return {"ok": True, "started": False, "running": True, "chunks": len(IMS_INDEX)}
    IMS_REINDEXING = True
    t = threading.Thread(target=_reindex_worker, daemon=True)
    t.start()
    return {"ok": True, "started": True, "running": True}

@app.get("/ims/_status")
def ims_status():
    files = {}
    for rec in IMS_INDEX:
        files[rec["relpath"]] = files.get(rec["relpath"], 0) + 1
    return {
        "running": IMS_REINDEXING,
        "chunks": len(IMS_INDEX),
        "files_indexed": len(files),
        "last_error": IMS_LAST_ERROR,
    }
# ===== end async patch =====

# ===== IMS diagnostics =====
@app.get("/ims/_where")
def ims_where():
    return {"ims_dir": IMS_DIR, "exists": os.path.isdir(IMS_DIR)}

@app.get("/ims/_ls")
def ims_list():
    if not os.path.isdir(IMS_DIR):
        return {"ims_dir": IMS_DIR, "exists": False, "files": []}
    out = []
    total = 0
    exts = {}
    # cap the listing to 300 items for safety
    for root, _, files in os.walk(IMS_DIR):
        for fn in files:
            total += 1
            ext = os.path.splitext(fn)[1].lower()
            exts[ext] = exts.get(ext, 0) + 1
            if len(out) < 300:
                rel = os.path.relpath(os.path.join(root, fn), IMS_DIR)
                out.append(rel)
    return {"ims_dir": IMS_DIR, "exists": True, "total_files": total, "by_ext": exts, "sample": out}
# ===== end diagnostics =====

# ---- Robust PDF extraction + graceful skip ----
def _safe_pdf_text(path, max_pages=None):
    """Extracts text from a PDF while tolerating broken xrefs, encrypted pages, etc."""
    out = []
    try:
        from pypdf import PdfReader
        reader = PdfReader(path, strict=False)
        n = len(reader.pages)
        limit = min(n, max_pages or n)
        for i in range(limit):
            try:
                pg = reader.pages[i]
                txt = pg.extract_text() or ""
                if txt.strip():
                    out.append(txt)
            except Exception as e:
                print(f"[ims][pdf][warn] {path} page {i} skipped: {e}")
                continue
        return "\n\n".join(out)
    except Exception as e:
        print(f"[ims][pdf][error] {path}: {e}")
        return ""

def _looks_scanned(path):
    # crude size heuristic: > 25MB often image-only scans; adjust if needed
    try:
        import os
        return os.path.getsize(path) > 25 * 1024 * 1024
    except Exception:
        return False

# Try to hook into existing extractor if present, else provide a default
try:
    # if project defines extract_text_from_file(), wrap it
    if 'extract_text_from_file' in globals():
        _old_etff = extract_text_from_file
        def extract_text_from_file(path):
            import os
            ext = os.path.splitext(path)[1].lower()
            if ext == ".pdf":
                if _looks_scanned(path):
                    print(f"[ims][pdf] skipping likely scanned PDF (large): {path}")
                    return ""
                return _safe_pdf_text(path)
            # fallback to original for other types
            return _old_etff(path)
        print("[ims] patched extract_text_from_file for robust PDF handling")
    else:
        # provide a minimal extractor if none existed
        import os
        from docx import Document as _Docx
        def extract_text_from_file(path):
            ext = os.path.splitext(path)[1].lower()
            if ext == ".pdf":
                if _looks_scanned(path):
                    print(f"[ims][pdf] skipping likely scanned PDF (large): {path}")
                    return ""
                return _safe_pdf_text(path)
            if ext in (".txt", ".md"):
                try:
                    return open(path, "r", errors="ignore").read()
                except Exception as e:
                    print(f"[ims][txt][error] {path}: {e}"); return ""
            if ext in (".docx",):
                try:
                    d = _Docx(path)
                    return "\n".join(p.text for p in d.paragraphs)
                except Exception as e:
                    print(f"[ims][docx][error] {path}: {e}"); return ""
            if ext in (".xlsx",):
                try:
                    import pandas as pd
                    texts=[]
                    x = pd.ExcelFile(path)
                    for s in x.sheet_names:
                        try:
                            df = x.parse(s, dtype=str).fillna("")
                            texts.append(df.to_csv(index=False))
                        except Exception as ee:
                            print(f"[ims][xlsx][warn] {path}:{s} {ee}")
                    return "\n\n".join(texts)
                except Exception as e:
                    print(f"[ims][xlsx][error] {path}: {e}"); return ""
            return ""  # unsupported
        print("[ims] provided default extract_text_from_file with robust PDF handling")
except Exception as e:
    print("[ims][patch] PDF extractor patch skipped:", e)
# ---- end patch ----
